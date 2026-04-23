"""
Document parser — extracts text from PDF, DOCX, and TXT files.

Why multiple parsers:
- PDFs are the most common business document format but hardest to parse
  (text can be in images, tables, multi-column layouts)
- DOCX is structured XML under the hood — easier to extract cleanly
- TXT is trivial but included for completeness

We use PyMuPDF (fitz) for PDFs because:
- Fastest pure-Python PDF parser
- Handles text extraction, table detection, and metadata
- Works without external dependencies (unlike pdfplumber/tabula)
"""

from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class ParsedDocument:
    """A parsed document with metadata."""
    filename: str
    text: str
    pages: int
    file_type: str
    metadata: dict = field(default_factory=dict)

    @property
    def total_chars(self) -> int:
        return len(self.text)


def parse_pdf(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    doc = fitz.open(str(path))

    pages_text = []
    for page in doc:
        text = page.get_text("text")
        if text.strip():
            pages_text.append(text)

    full_text = "\n\n".join(pages_text)

    result = ParsedDocument(
        filename=path.name,
        text=full_text,
        pages=len(doc),
        file_type="pdf",
        metadata={
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
        },
    )
    doc.close()
    return result


def parse_docx(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    doc = DocxDocument(str(path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)

    return ParsedDocument(
        filename=path.name,
        text=full_text,
        pages=1,  # DOCX doesn't have page concept
        file_type="docx",
    )


def parse_txt(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    text = path.read_text(encoding="utf-8", errors="replace")

    return ParsedDocument(
        filename=path.name,
        text=text,
        pages=1,
        file_type="txt",
    )


def parse_document(file_path: str) -> ParsedDocument:
    """Auto-detect file type and parse accordingly."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".txt": parse_txt,
        ".md": parse_txt,
    }

    parser = parsers.get(suffix)
    if parser is None:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported: {list(parsers.keys())}"
        )

    result = parser(file_path)
    print(f"  Parsed {result.filename}: {result.total_chars} chars, {result.pages} pages")
    return result
